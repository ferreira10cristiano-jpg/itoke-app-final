"""
Script to generate the iToke Tax Planning document for the accountant.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(30, 41, 59)

# ===================== COVER =====================
doc.add_paragraph()
title = doc.add_heading('', level=0)
run = title.add_run('iToke — Plano de Acao Fiscal e Tributario')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(16, 185, 129)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Documento para Consultoria Contabil\nVersao 1.0 — Abril 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONFIDENCIAL — USO INTERNO')
run.bold = True
run.font.color.rgb = RGBColor(239, 68, 68)

doc.add_page_break()

# ===================== INDEX =====================
doc.add_heading('INDICE', level=1)
items = [
    '1. Descricao do Negocio e Modelo de Receita',
    '2. Fluxo Financeiro Detalhado',
    '3. Analise de CNAE Recomendado',
    '4. Comparativo de Regimes Tributarios',
    '5. Obrigacoes Fiscais e Documentais',
    '6. Riscos Fiscais Identificados',
    '7. Plano de Acao — Curto, Medio e Longo Prazo',
    '8. Impacto da Reforma Tributaria (2026-2027)',
    '9. Checklist para o Contador',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ===================== 1. DESCRICAO =====================
doc.add_heading('1. DESCRICAO DO NEGOCIO E MODELO DE RECEITA', level=1)
doc.add_paragraph()
doc.add_paragraph('O iToke e uma plataforma digital (app Android/iOS + web) que conecta estabelecimentos comerciais a clientes por meio de um sistema de tokens e ofertas com desconto.')
doc.add_paragraph()

doc.add_heading('Como funciona:', level=3)
doc.add_paragraph('1. O ESTABELECIMENTO compra pacotes de tokens na plataforma (ex: 10 tokens por R$ 9,90)', style='List Bullet')
doc.add_paragraph('2. O ESTABELECIMENTO cria ofertas com desconto e aloca tokens nelas', style='List Bullet')
doc.add_paragraph('3. O CLIENTE encontra a oferta, gera um QR Code (gasta 1 token ou credito) e vai ao local', style='List Bullet')
doc.add_paragraph('4. O ESTABELECIMENTO valida o QR Code e finaliza a venda', style='List Bullet')
doc.add_paragraph('5. O CLIENTE que indicou outros clientes recebe creditos (comissao de rede)', style='List Bullet')
doc.add_paragraph('6. O ESTABELECIMENTO acumula creditos dos clientes e pode solicitar saque via PIX', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Modelo de Receita do iToke:', level=3)
doc.add_paragraph()

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Fonte de Receita', 'Exemplo', 'Valor iToke']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('Venda de tokens aos estabelecimentos', '10 tokens por R$ 9,90', 'R$ 6,90 (margem iToke)'),
    ('Venda de tokens aos clientes', 'Pacotes na loja do app', 'R$ 6,90 (margem iToke)'),
    ('Repasse de creditos ao estabelecimento', 'Quando cliente usa credito no QR', 'R$ 3,00 por token (REPASSE, nao receita)'),
    ('Comissao de representantes', 'R$ 1,00 por transacao', 'Custo operacional'),
]
for i, (a, b, c) in enumerate(data):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b
    table.rows[i+1].cells[2].text = c

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('IMPORTANTE: ')
run.bold = True
run.font.color.rgb = RGBColor(239, 68, 68)
p.add_run('O valor de R$ 3,00 por token que vai para o estabelecimento e REPASSE, nao receita do iToke. O iToke atua como intermediador. A receita real do iToke e apenas a margem (R$ 6,90 no exemplo de pacote de R$ 9,90).')

doc.add_page_break()

# ===================== 2. FLUXO FINANCEIRO =====================
doc.add_heading('2. FLUXO FINANCEIRO DETALHADO', level=1)
doc.add_paragraph()

doc.add_heading('Fluxo 1 — Venda de Tokens (Entrada de Receita)', level=3)
doc.add_paragraph('Cliente ou Estabelecimento compra pacote via Stripe (cartao de credito/debito)')
doc.add_paragraph()

table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
flow1 = [
    ('Valor pago pelo comprador', 'R$ 9,90'),
    ('Taxa Stripe (~3,99% + R$ 0,39)', '-R$ 0,79'),
    ('Valor liquido recebido pelo iToke', 'R$ 9,11'),
    ('Margem iToke (receita tributavel)', 'R$ 6,90'),
    ('Valor repassavel ao estabelecimento', 'R$ 3,00'),
    ('Custo operacional (servidor, etc.)', 'Variavel'),
]
for i, (k, v) in enumerate(flow1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_heading('Fluxo 2 — Resgate de Creditos (Saida — Saque do Estabelecimento)', level=3)
doc.add_paragraph('Quando o estabelecimento solicita saque dos creditos acumulados:')
doc.add_paragraph('1. Estabelecimento solicita saque via dashboard', style='List Bullet')
doc.add_paragraph('2. Admin do iToke aprova o saque', style='List Bullet')
doc.add_paragraph('3. iToke realiza transferencia PIX para o estabelecimento', style='List Bullet')
doc.add_paragraph('4. Registro da transacao com data, valor e comprovante', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Fluxo 3 — Comissao de Representantes (Saida)', level=3)
doc.add_paragraph('Representante PJ ganha R$ 1,00 por transacao dos seus indicados (editavel pelo admin).')
doc.add_paragraph('Representante solicita saque via PIX, aprovado pelo admin.')

doc.add_page_break()

# ===================== 3. CNAE =====================
doc.add_heading('3. ANALISE DE CNAE RECOMENDADO', level=1)
doc.add_paragraph()

doc.add_heading('CNAE Principal Recomendado:', level=3)
p = doc.add_paragraph()
run = p.add_run('7490-1/04 — Atividades de intermediacao e agenciamento de servicos e negocios em geral, exceto imobiliarios')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('Justificativa:', style='List Bullet')
doc.add_paragraph('O iToke conecta estabelecimentos comerciais a clientes, cobrando taxa/comissao por transacao', style='List Bullet 2')
doc.add_paragraph('Nao vende produtos fisicos, nao armazena mercadorias', style='List Bullet 2')
doc.add_paragraph('Receita baseada em intermediacao digital', style='List Bullet 2')
doc.add_paragraph('CNAE reconhecido pelo IBGE para marketplaces e plataformas de intermediacao', style='List Bullet 2')

doc.add_paragraph()
doc.add_heading('CNAEs Secundarios Sugeridos:', level=3)

table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'CNAE'
table.rows[0].cells[1].text = 'Descricao'
cnaes = [
    ('6209-1/00', 'Suporte tecnico, manutencao e outros servicos em TI'),
    ('6311-9/00', 'Tratamento de dados, provedores de servicos de aplicacao e hospedagem'),
    ('7319-0/02', 'Promocao de vendas (para a parte de marketing e ofertas)'),
]
for i, (c, d) in enumerate(cnaes):
    table.rows[i+1].cells[0].text = c
    table.rows[i+1].cells[1].text = d

doc.add_page_break()

# ===================== 4. REGIMES TRIBUTARIOS =====================
doc.add_heading('4. COMPARATIVO DE REGIMES TRIBUTARIOS', level=1)
doc.add_paragraph()

doc.add_heading('Cenario: Faturamento estimado nos primeiros 12 meses', level=3)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Cenario'
table.rows[0].cells[1].text = 'Faturamento Anual'
table.rows[0].cells[2].text = 'Regime Recomendado'
scenarios = [
    ('Inicial (validacao)', 'Ate R$ 180.000', 'MEI (se possivel) ou Simples Nacional'),
    ('Crescimento', 'R$ 180.000 — R$ 500.000', 'Simples Nacional (Anexo III ou V)'),
    ('Escala', 'R$ 500.000 — R$ 4.800.000', 'Simples Nacional ou Lucro Presumido (simular)'),
]
for i, (a, b, c) in enumerate(scenarios):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b
    table.rows[i+1].cells[2].text = c

doc.add_paragraph()
doc.add_heading('Simples Nacional — Anexo V (Servicos de intermediacao)', level=3)

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Faixa'
table.rows[0].cells[1].text = 'Faturamento 12 meses'
table.rows[0].cells[2].text = 'Aliquota'
faixas = [
    ('1a', 'Ate R$ 180.000', '15,50%'),
    ('2a', 'R$ 180.000 — R$ 360.000', '18,00%'),
    ('3a', 'R$ 360.000 — R$ 720.000', '19,50%'),
    ('4a', 'R$ 720.000 — R$ 1.800.000', '20,50%'),
    ('5a', 'R$ 1.800.000 — R$ 3.600.000', '23,00%'),
    ('6a', 'R$ 3.600.000 — R$ 4.800.000', '30,50%'),
]
for i, (a, b, c) in enumerate(faixas):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b
    table.rows[i+1].cells[2].text = c

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('FATOR R: ')
run.bold = True
p.add_run('Se a folha de pagamento (incluindo pro-labore) for >= 28% do faturamento bruto dos ultimos 12 meses, o iToke migra do Anexo V para o Anexo III (aliquotas menores, a partir de 6%). Estrategia: manter pro-labore adequado para atingir o Fator R.')

doc.add_paragraph()
doc.add_heading('Lucro Presumido (para comparacao)', level=3)

table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Imposto'
table.rows[0].cells[1].text = 'Aliquota sobre faturamento'
lp = [
    ('IRPJ', '4,80% (15% sobre presuncao de 32%)'),
    ('CSLL', '2,88% (9% sobre presuncao de 32%)'),
    ('PIS', '0,65%'),
    ('COFINS', '3,00%'),
    ('ISS', '2,00% — 5,00% (varia por municipio)'),
]
for i, (a, b) in enumerate(lp):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Total Lucro Presumido: ~13,33% a 16,33% do faturamento')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('RECOMENDACAO INICIAL: ')
run.bold = True
run.font.color.rgb = RGBColor(16, 185, 129)
p.add_run('Simples Nacional com estrategia de Fator R (Anexo III). Simular com o contador a cada trimestre se vale migrar para Lucro Presumido conforme o faturamento cresce.')

doc.add_page_break()

# ===================== 5. OBRIGACOES FISCAIS =====================
doc.add_heading('5. OBRIGACOES FISCAIS E DOCUMENTAIS', level=1)
doc.add_paragraph()

doc.add_heading('5.1 Notas Fiscais', level=3)
doc.add_paragraph('O iToke deve emitir NFS-e (Nota Fiscal de Servico Eletronica) sobre a RECEITA DE INTERMEDIACAO (margem), NAO sobre o valor total da transacao.')
doc.add_paragraph()
doc.add_paragraph('Exemplo: Pacote vendido por R$ 9,90. NFS-e emitida sobre R$ 6,90 (servico de intermediacao digital). Os R$ 3,00 restantes sao repasse ao estabelecimento.', style='List Bullet')

doc.add_paragraph()
doc.add_heading('5.2 Documentacao de Repasses', level=3)
doc.add_paragraph('Para cada saque realizado pelo estabelecimento, manter:', style='List Bullet')
doc.add_paragraph('Relatorio de repasse com data, valor, destinatario e comprovante PIX', style='List Bullet 2')
doc.add_paragraph('Referencia ao contrato de intermediacao', style='List Bullet 2')
doc.add_paragraph('ID da transacao no sistema', style='List Bullet 2')

doc.add_paragraph()
doc.add_heading('5.3 Contratos Obrigatorios', level=3)
doc.add_paragraph('Contrato de Intermediacao com Estabelecimentos — define que o iToke e intermediador, nao vendedor', style='List Bullet')
doc.add_paragraph('Contrato de Representacao Comercial PJ — para representantes (ja implementado no app com assinatura digital)', style='List Bullet')
doc.add_paragraph('Termos de Uso e Politica de Privacidade — para clientes (ja implementado)', style='List Bullet')

doc.add_paragraph()
doc.add_heading('5.4 LGPD — Protecao de Dados', level=3)
doc.add_paragraph('CPFs de clientes nos recibos de estabelecimentos devem ser MASCARADOS (ex: 123.***.***-01)', style='List Bullet')
doc.add_paragraph('Dados pessoais armazenados com consentimento e para finalidade especifica', style='List Bullet')
doc.add_paragraph('Politica de retencao e exclusao de dados definida', style='List Bullet')

doc.add_page_break()

# ===================== 6. RISCOS =====================
doc.add_heading('6. RISCOS FISCAIS IDENTIFICADOS', level=1)
doc.add_paragraph()

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Risco'
table.rows[0].cells[1].text = 'Impacto'
table.rows[0].cells[2].text = 'Mitigacao'
risks = [
    ('Receita Federal considerar valor total como receita', 'Pagar imposto sobre 100% e nao apenas a margem', 'Contrato de intermediacao + NFS-e apenas sobre margem + relatorio de repasses'),
    ('Ausencia de NFS-e nos repasses', 'Autuacao e multa', 'Emitir NFS-e sobre servico de intermediacao'),
    ('CPF de clientes expostos em recibos', 'Multa LGPD (ate 2% do faturamento, max R$ 50M)', 'Mascarar CPFs: 123.***.***-01'),
    ('Representante configurado como "empregado"', 'Vinculo empregaticio + FGTS + INSS', 'Contrato PJ explicito, sem subordinacao, pagamento por resultado'),
    ('Falta de separacao contabil entre receita e repasse', 'Bitributacao', 'Plano de contas separado: Receita vs Repasse'),
]
for i, (a, b, c) in enumerate(risks):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b
    table.rows[i+1].cells[2].text = c

doc.add_page_break()

# ===================== 7. PLANO DE ACAO =====================
doc.add_heading('7. PLANO DE ACAO', level=1)
doc.add_paragraph()

doc.add_heading('Curto Prazo (Imediato — Proximo mes)', level=2)
doc.add_paragraph('Abrir CNPJ com CNAE 7490-1/04 (se ainda nao tem)', style='List Bullet')
doc.add_paragraph('Optar pelo Simples Nacional', style='List Bullet')
doc.add_paragraph('Definir pro-labore adequado para atingir Fator R >= 28%', style='List Bullet')
doc.add_paragraph('Elaborar contrato de intermediacao com estabelecimentos', style='List Bullet')
doc.add_paragraph('Configurar emissao de NFS-e sobre servico de intermediacao', style='List Bullet')
doc.add_paragraph('Mascarar CPFs em todos os recibos e relatorios (JA FEITO no app)', style='List Bullet')
doc.add_paragraph('Separar contas bancarias: conta operacional (repasses) vs conta de receita', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Medio Prazo (3 a 6 meses)', level=2)
doc.add_paragraph('Implementar Stripe Connect para split payment automatico (separa comissao do repasse)', style='List Bullet')
doc.add_paragraph('Criar plano de contas contabil: Receita de Intermediacao vs Repasses a Terceiros', style='List Bullet')
doc.add_paragraph('Simular trimestralmente: Simples Nacional vs Lucro Presumido', style='List Bullet')
doc.add_paragraph('Contratar seguro de responsabilidade civil digital', style='List Bullet')
doc.add_paragraph('Implementar emissao automatica de NFS-e integrada ao app', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Longo Prazo (2026-2027 — Reforma Tributaria)', level=2)
doc.add_paragraph('Adaptar sistema para split payment obrigatorio (IBS + CBS)', style='List Bullet')
doc.add_paragraph('Integrar com sistema de NF-e com destaque de IBS/CBS', style='List Bullet')
doc.add_paragraph('Monitorar Notas Tecnicas da Receita Federal', style='List Bullet')
doc.add_paragraph('Reavaliar regime tributario apos reforma', style='List Bullet')

doc.add_page_break()

# ===================== 8. REFORMA TRIBUTARIA =====================
doc.add_heading('8. IMPACTO DA REFORMA TRIBUTARIA (2026-2027)', level=1)
doc.add_paragraph()

doc.add_paragraph('A Reforma Tributaria (EC 132/2023 + LC 214/2025) impacta diretamente plataformas digitais:')
doc.add_paragraph()
doc.add_paragraph('2026: Periodo de testes. NF-e com destaque de IBS/CBS (sem cobranca real). Plataformas devem adaptar sistemas.', style='List Bullet')
doc.add_paragraph('2027: Split payment obrigatorio. Bancos e arranjos de pagamento (Pix, cartao) dividem automaticamente: valor do vendedor + impostos + comissao da plataforma.', style='List Bullet')
doc.add_paragraph('Responsabilidade solidaria: Marketplaces respondem pelo recolhimento de IBS/CBS dos vendedores parceiros.', style='List Bullet')
doc.add_paragraph('Aliquota estimada IVA: ~26,5% (unificando PIS/COFINS/ICMS/ISS).', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('ACAO NECESSARIA: ')
run.bold = True
p.add_run('Iniciar adaptacao dos sistemas em 2026 (periodo de testes) para estar em conformidade em 2027.')

doc.add_page_break()

# ===================== 9. CHECKLIST =====================
doc.add_heading('9. CHECKLIST PARA O CONTADOR', level=1)
doc.add_paragraph()
doc.add_paragraph('Use este checklist para orientar a consultoria contabil:')
doc.add_paragraph()

checks = [
    'Confirmar CNAE 7490-1/04 como principal (intermediacao)',
    'Enquadrar no Simples Nacional (Anexo V, com estrategia Fator R para Anexo III)',
    'Definir pro-labore minimo para atingir Fator R >= 28%',
    'Elaborar modelo de NFS-e para servico de intermediacao digital',
    'Criar contrato de intermediacao entre iToke e Estabelecimentos',
    'Definir plano de contas: Receita de Servico vs Repasses a Terceiros',
    'Validar que repasses NAO sao tributados como receita',
    'Orientar sobre obrigacoes acessorias (DEFIS, DAS, etc.)',
    'Simular carga tributaria: Simples vs Lucro Presumido',
    'Avaliar necessidade de inscricao municipal para ISS',
    'Revisar conformidade LGPD nos dados armazenados',
    'Planejar adaptacao para Reforma Tributaria 2027 (split payment)',
    'Avaliar necessidade de conta escrow ou conta intermediaria para repasses',
]

for c in checks:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento gerado pela plataforma iToke — Abril 2026')
run.italic = True
run.font.color.rgb = RGBColor(148, 163, 184)

output_path = '/app/frontend/iToke_Plano_Fiscal_Contador.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
