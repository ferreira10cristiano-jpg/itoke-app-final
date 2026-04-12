"""
Script to generate the iToke PowerPoint presentation prompt in Word format.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(30, 41, 59)

# ===================== COVER =====================
doc.add_paragraph()
title = doc.add_heading('', level=0)
run = title.add_run('iToke')
run.font.size = Pt(44)
run.font.color.rgb = RGBColor(16, 185, 129)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Ofertas que saem de Graca')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(100, 116, 139)
run.italic = True

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Apresentacao Comercial\nPara Estabelecimentos e Representantes')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(51, 65, 85)

doc.add_paragraph()
doc.add_paragraph()

# ===================== INSTRUCTIONS =====================
doc.add_heading('INSTRUCOES PARA CRIAR O POWERPOINT', level=1)
p = doc.add_paragraph()
p.add_run('Este documento e um roteiro completo para voce criar uma apresentacao profissional em PowerPoint. ').bold = False
p.add_run('Cada secao abaixo corresponde a um SLIDE. ').bold = True
p.add_run('Use as capturas de tela do app disponibilizadas junto com este arquivo.')

doc.add_paragraph()
tips = doc.add_paragraph()
tips.add_run('Dicas de Design:').bold = True
doc.add_paragraph('Fundo escuro (#0F172A ou #0B0F1A) para combinar com o app', style='List Bullet')
doc.add_paragraph('Cor principal: Verde (#10B981) para destaques e titulos', style='List Bullet')
doc.add_paragraph('Cor secundaria: Azul (#3B82F6) para icones e graficos', style='List Bullet')
doc.add_paragraph('Fonte sugerida: Montserrat ou Poppins (moderna e limpa)', style='List Bullet')
doc.add_paragraph('Inserir capturas de tela do app em mockups de celular', style='List Bullet')

doc.add_page_break()

# ===================== SLIDE 1 =====================
doc.add_heading('SLIDE 1 — CAPA', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('iToke — Ofertas que saem de Graca')
p = doc.add_paragraph()
p.add_run('Subtitulo: ').bold = True
p.add_run('A plataforma que conecta clientes a ofertas incriveis e gera receita para seu negocio')
p = doc.add_paragraph()
p.add_run('Imagem: ').bold = True
p.add_run('Logo do iToke centralizado + mockup do celular com a tela inicial (tela_inicial.png)')
p = doc.add_paragraph()
p.add_run('Rodape: ').bold = True
p.add_run('www.itoke.com.br | Disponivel no Google Play')

doc.add_paragraph()
doc.add_paragraph('[INSERIR: Captura da tela inicial do app - tela_inicial.png]', style='Intense Quote')

doc.add_page_break()

# ===================== SLIDE 2 =====================
doc.add_heading('SLIDE 2 — O PROBLEMA', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('O Desafio dos Estabelecimentos')
doc.add_paragraph()
doc.add_paragraph('Atrair novos clientes esta cada vez mais caro', style='List Bullet')
doc.add_paragraph('Plataformas de delivery cobram taxas de 25-30%', style='List Bullet')
doc.add_paragraph('Clientes nao retornam sem incentivo', style='List Bullet')
doc.add_paragraph('Marketing digital e complexo e caro para pequenos negocios', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Visual: ').bold = True
p.add_run('Icones representando cada problema (dinheiro voando, grafico caindo, pessoa saindo)')

doc.add_page_break()

# ===================== SLIDE 3 =====================
doc.add_heading('SLIDE 3 — A SOLUCAO iToke', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Como o iToke Funciona')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Texto central: ').bold = True
p.add_run('"Uma plataforma simples onde estabelecimentos criam ofertas com desconto, clientes resgatam usando QR Codes, e todos ganham."')
doc.add_paragraph()
doc.add_heading('Fluxo Visual (3 passos):', level=3)
doc.add_paragraph('1. ESTABELECIMENTO compra tokens e cria ofertas com desconto', style='List Bullet')
doc.add_paragraph('2. CLIENTE encontra a oferta, gera QR Code e vai ao local', style='List Bullet')
doc.add_paragraph('3. ESTABELECIMENTO valida o QR e finaliza a venda', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Visual: ').bold = True
p.add_run('Setas conectando icones: Loja -> Celular -> QR Code -> Venda')

doc.add_page_break()

# ===================== SLIDE 4 =====================
doc.add_heading('SLIDE 4 — VISAO DO CLIENTE', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('A Experiencia do Cliente')
doc.add_paragraph()
doc.add_heading('Funcionalidades:', level=3)
doc.add_paragraph('Explorar ofertas por cidade, bairro e categoria', style='List Bullet')
doc.add_paragraph('Ver desconto real (ate 60% OFF)', style='List Bullet')
doc.add_paragraph('Gerar QR Code para resgatar a oferta', style='List Bullet')
doc.add_paragraph('Ganhar creditos indicando amigos', style='List Bullet')
doc.add_paragraph('Comprar pacotes de tokens', style='List Bullet')
doc.add_paragraph('Historico de compras com recibo PDF', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Imagens: ').bold = True
p.add_run('Capturas da tela Home (ofertas) e tela de Comprar Tokens lado a lado')
doc.add_paragraph()
doc.add_paragraph('[INSERIR: screen_home.png + screen_buytokens.png em mockups de celular]', style='Intense Quote')

doc.add_page_break()

# ===================== SLIDE 5 =====================
doc.add_heading('SLIDE 5 — VISAO DO ESTABELECIMENTO', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Dashboard do Estabelecimento')
doc.add_paragraph()
doc.add_heading('O que o estabelecimento pode fazer:', level=3)
doc.add_paragraph('Dashboard completo com metricas (visualizacoes, QRs, vendas)', style='List Bullet')
doc.add_paragraph('Criar e gerenciar ofertas com fotos', style='List Bullet')
doc.add_paragraph('Comprar tokens para ativar ofertas', style='List Bullet')
doc.add_paragraph('Validar QR Codes dos clientes', style='List Bullet')
doc.add_paragraph('Relatorio fiscal em PDF', style='List Bullet')
doc.add_paragraph('Solicitar saques dos creditos recebidos via PIX', style='List Bullet')
doc.add_paragraph('Centro de aprendizado com videos', style='List Bullet')
doc.add_paragraph()

# Stats highlight
table = doc.add_table(rows=2, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Visualizacoes', 'QR Gerados', 'Vendas', 'Ofertas']
values = ['3.176', '60', '22', '11']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[1].cells[i].text = values[i]

doc.add_paragraph()
doc.add_paragraph('[INSERIR: screen_establishment.png em mockup de celular]', style='Intense Quote')

doc.add_page_break()

# ===================== SLIDE 6 =====================
doc.add_heading('SLIDE 6 — COMO O ESTABELECIMENTO GANHA', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Modelo de Receita para o Estabelecimento')
doc.add_paragraph()
doc.add_heading('Exemplo pratico:', level=3)
doc.add_paragraph()

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ('Preco original do produto', 'R$ 50,00'),
    ('Desconto ofertado (40%)', 'R$ 30,00'),
    ('Cliente paga COM credito', 'R$ 25,00 (dinheiro) + R$ 5,00 (credito)'),
    ('Estabelecimento recebe', 'R$ 30,00 total'),
    ('Custo do token usado', 'R$ 2,00'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resultado: ').bold = True
p.add_run('O estabelecimento atrai um NOVO CLIENTE, faz uma venda de R$ 30 investindo apenas R$ 2 em token. Custo de aquisicao: R$ 2,00 por cliente!')

doc.add_page_break()

# ===================== SLIDE 7 =====================
doc.add_heading('SLIDE 7 — SISTEMA DE REPRESENTANTES', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Seja um Representante Comercial iToke')
doc.add_paragraph()
doc.add_heading('Como funciona:', level=3)
doc.add_paragraph('Representante e cadastrado com CNPJ pela equipe iToke', style='List Bullet')
doc.add_paragraph('Recebe um codigo exclusivo de indicacao', style='List Bullet')
doc.add_paragraph('Indica novos estabelecimentos e clientes', style='List Bullet')
doc.add_paragraph('Ganha comissao de R$ 1,00 por CADA transacao dos seus indicados', style='List Bullet')
doc.add_paragraph('Vinculo de 12 meses com seus indicados', style='List Bullet')
doc.add_paragraph('Recebe tokens gratuitos para oferecer a novos estabelecimentos', style='List Bullet')
doc.add_paragraph('Dashboard proprio para acompanhar tudo', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Destaque: ').bold = True
p.add_run('"NAO e multinivel. Comissao direta e transparente."')
doc.add_paragraph()
doc.add_paragraph('[INSERIR: screen_rep.png em mockup de celular]', style='Intense Quote')

doc.add_page_break()

# ===================== SLIDE 8 =====================
doc.add_heading('SLIDE 8 — SIMULACAO DE GANHOS DO REPRESENTANTE', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Quanto um Representante Pode Ganhar?')
doc.add_paragraph()

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Cenario', 'Estabelecimentos', 'Transacoes/mes', 'Ganho Mensal']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('Iniciante', '5', '50', 'R$ 50,00'),
    ('Ativo', '15', '200', 'R$ 200,00'),
    ('Top', '30', '500', 'R$ 500,00'),
    ('Estrela', '50', '1.000', 'R$ 1.000,00'),
]
for i, (a, b, c, d) in enumerate(data):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b
    table.rows[i+1].cells[2].text = c
    table.rows[i+1].cells[3].text = d

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Nota: ').bold = True
p.add_run('Comissao de R$ 1,00 por transacao. Valores ilustrativos baseados em media de uso.')

doc.add_page_break()

# ===================== SLIDE 9 =====================
doc.add_heading('SLIDE 9 — DASHBOARD DO REPRESENTANTE', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Controle Total na Palma da Mao')
doc.add_paragraph()
doc.add_heading('O representante acompanha:', level=3)
doc.add_paragraph('Numero de clientes e estabelecimentos indicados', style='List Bullet')
doc.add_paragraph('Saldo de comissoes acumuladas', style='List Bullet')
doc.add_paragraph('Total ja ganho', style='List Bullet')
doc.add_paragraph('Tokens gratuitos (alocados, usados, restantes)', style='List Bullet')
doc.add_paragraph('Historico de comissoes detalhado', style='List Bullet')
doc.add_paragraph('Solicitacao de saque via PIX', style='List Bullet')
doc.add_paragraph('Contrato digital assinado', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Visual: ').bold = True
p.add_run('Screenshot do dashboard do representante com as metricas e abas')
doc.add_paragraph()
doc.add_paragraph('[INSERIR: screen_rep.png mostrando stats, tokens gratuitos e abas]', style='Intense Quote')

doc.add_page_break()

# ===================== SLIDE 10 =====================
doc.add_heading('SLIDE 10 — SEGURANCA E ANTI-FRAUDE', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Plataforma Segura e Confiavel')
doc.add_paragraph()
doc.add_paragraph('Pagamento seguro via Stripe (lider mundial)', style='List Bullet')
doc.add_paragraph('Validacao de CPF com algoritmo oficial', style='List Bullet')
doc.add_paragraph('Validacao de CNPJ para representantes', style='List Bullet')
doc.add_paragraph('Sistema anti-fraude com rate limiting', style='List Bullet')
doc.add_paragraph('Deteccao de auto-indicacao (CNPJ cruzado)', style='List Bullet')
doc.add_paragraph('Alertas em tempo real para a equipe', style='List Bullet')
doc.add_paragraph('Contrato digital com registro de IP e timestamp', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Visual: ').bold = True
p.add_run('Icone de escudo/cadeado verde com checkmarks ao lado de cada item')

doc.add_page_break()

# ===================== SLIDE 11 =====================
doc.add_heading('SLIDE 11 — PROXIMOS PASSOS', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Como Comecar?')
doc.add_paragraph()
doc.add_heading('Para Estabelecimentos:', level=3)
doc.add_paragraph('1. Baixe o app iToke na Google Play', style='List Bullet')
doc.add_paragraph('2. Selecione "Sou Estabelecimento"', style='List Bullet')
doc.add_paragraph('3. Preencha seus dados e CNPJ', style='List Bullet')
doc.add_paragraph('4. Compre seu primeiro pacote de tokens', style='List Bullet')
doc.add_paragraph('5. Crie sua primeira oferta e comece a vender!', style='List Bullet')
doc.add_paragraph()
doc.add_heading('Para Representantes:', level=3)
doc.add_paragraph('1. Entre em contato com a equipe iToke', style='List Bullet')
doc.add_paragraph('2. Envie seus documentos (CNPJ, RG, Contrato Social)', style='List Bullet')
doc.add_paragraph('3. Aceite o contrato digital', style='List Bullet')
doc.add_paragraph('4. Receba seu codigo e comece a indicar!', style='List Bullet')

doc.add_page_break()

# ===================== SLIDE 12 =====================
doc.add_heading('SLIDE 12 — CONTATO (SLIDE FINAL)', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Titulo: ').bold = True
p.add_run('Vamos Crescer Juntos!')
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('iToke — Ofertas que saem de Graca').bold = True
p.font_size = Pt(16)
doc.add_paragraph()
doc.add_paragraph('Site: www.itoke.com.br')
doc.add_paragraph('App: Disponivel na Google Play')
doc.add_paragraph('Telefone: (61) 99165-3968')
doc.add_paragraph('E-mail: contato@itoke.com.br')
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Transforme seus clientes em fas. Transforme suas vendas em lucro."')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(16, 185, 129)

doc.add_page_break()

# ===================== APPENDIX =====================
doc.add_heading('ANEXO — LISTA DE CAPTURAS DE TELA', level=1)
doc.add_paragraph()
doc.add_paragraph('As seguintes capturas de tela do app estao disponibilizadas para uso na apresentacao:')
doc.add_paragraph()

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Arquivo'
table.rows[0].cells[1].text = 'Descricao'
table.rows[0].cells[2].text = 'Usar no Slide'
files = [
    ('tela_inicial.png', 'Tela de boas-vindas (Sou Cliente / Sou Estabelecimento)', 'Slide 1'),
    ('screen_home.png', 'Home do cliente com ofertas, categorias e descontos', 'Slide 4'),
    ('screen_buytokens.png', 'Tela de compra de tokens (pacotes e precos)', 'Slide 4'),
    ('screen_establishment.png', 'Dashboard do estabelecimento (metricas e saldo)', 'Slides 5-6'),
    ('screen_rep.png', 'Dashboard do representante (stats, tokens, abas)', 'Slides 7-9'),
    ('screen_wallet.png', 'Tela de creditos e rede de indicacao', 'Slide 4 (extra)'),
]
for i, (f, d, s) in enumerate(files):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = d
    table.rows[i+1].cells[2].text = s

doc.add_paragraph()
doc.add_heading('Dica Final:', level=3)
doc.add_paragraph('Para um visual profissional, use mockups de celular (smartphone frames) para inserir as capturas de tela. Sites gratuitos como mockuphone.com ou smartmockups.com permitem gerar mockups rapidamente.')

# Save
output_path = '/app/frontend/iToke_Apresentacao_Comercial.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
