from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== CAPA =====
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('iToke')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(16, 185, 129)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Plano de Implantação e Roadmap')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(15, 23, 42)

slogan = doc.add_paragraph()
slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = slogan.add_run('"Ofertas que saem de Graça"')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph('')
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Abril 2026 - Versão 1.0')
run.font.size = Pt(12)

doc.add_page_break()

# ===== ÍNDICE =====
doc.add_heading('Índice', level=1)
items = [
    '1. Visão Geral do Projeto',
    '2. Status Atual (O que já foi feito)',
    '3. Infraestrutura de Produção',
    '4. Fluxo de Atualizações',
    '5. Planejamento de Implantação - Fase 1 (Pré-Lançamento)',
    '6. Planejamento de Implantação - Fase 2 (Lançamento)',
    '7. Planejamento de Implantação - Fase 3 (Pós-Lançamento)',
    '8. Cronograma Detalhado',
    '9. Custos Operacionais',
    '10. Credenciais e Acessos',
    '11. Checklist de Lançamento',
    '12. Contatos e Suporte',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ===== 1. VISÃO GERAL =====
doc.add_heading('1. Visão Geral do Projeto', level=1)
doc.add_paragraph(
    'O iToke é uma plataforma de intermediação de ofertas onde estabelecimentos '
    'criam promoções e clientes resgatam usando tokens digitais. O modelo de negócio '
    'é baseado em indicações em 3 níveis, onde clientes ganham créditos por indicar '
    'novos usuários.'
)

doc.add_heading('Público-alvo:', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Tipo', 'Quem é', 'O que faz']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['Cliente', 'Consumidor final', 'Compra tokens, resgata ofertas, indica amigos'],
    ['Estabelecimento', 'Lojista/Comerciante', 'Cria ofertas, valida QR codes, compra tokens'],
    ['Admin', 'Gestor da plataforma', 'Gerencia tudo: usuários, ofertas, comissões, termos legais'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i+1].cells[j].text = cell

doc.add_page_break()

# ===== 2. STATUS ATUAL =====
doc.add_heading('2. Status Atual (O que já foi feito)', level=1)

doc.add_heading('Funcionalidades 100% Implementadas:', level=2)
features = [
    ['Sistema de Autenticação', 'Login/registro por email, gestão de sessões, CPF obrigatório'],
    ['Painel Admin', 'Dashboard completo, gestão de usuários, ofertas, comissões, termos legais, config da loja'],
    ['Painel Estabelecimento', 'Criar/editar ofertas, validar QR codes, comprar tokens, relatório fiscal com PDF'],
    ['App Cliente', 'Explorar ofertas, gerar QR code, resgatar ofertas, indicar amigos, comprar tokens'],
    ['FAQ', 'Perguntas frequentes com vídeos do YouTube incorporados'],
    ['Relatório Fiscal', 'Dashboard fiscal com exportação em PDF para estabelecimentos'],
    ['Documentos Legais', 'Termos de Uso, Política de Privacidade, editáveis pelo Admin'],
    ['Integração Stripe', 'Pagamento real de tokens via checkout Stripe (modo teste ativo)'],
    ['Deploy Backend', 'Servidor rodando no Railway com MongoDB Atlas na nuvem'],
    ['Build Android', 'APK de preview + AAB de produção gerados via Expo EAS'],
    ['Play Store', 'App publicado como Teste Interno, aguardando revisão do Google'],
]

table = doc.add_table(rows=len(features)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Funcionalidade'
table.rows[0].cells[1].text = 'Detalhes'
for p in table.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for i, (feat, detail) in enumerate(features):
    table.rows[i+1].cells[0].text = feat
    table.rows[i+1].cells[1].text = detail

doc.add_page_break()

# ===== 3. INFRAESTRUTURA =====
doc.add_heading('3. Infraestrutura de Produção', level=1)

doc.add_heading('Serviços Ativos:', level=2)
infra = [
    ['Serviço', 'Provedor', 'URL/Endereço', 'Status'],
    ['Backend API', 'Railway', 'https://itoke-app-final-production.up.railway.app', 'ONLINE'],
    ['Banco de Dados', 'MongoDB Atlas', 'cluster0.uxjrdiy.mongodb.net', 'ONLINE'],
    ['Código Fonte', 'GitHub', 'github.com/ferreira10cristiano-jpg/itoke-app-final', 'OK'],
    ['Build Mobile', 'Expo EAS', 'expo.dev/accounts/itokecris/projects/itoke', 'OK'],
    ['Loja de Apps', 'Google Play Store', 'play.google.com/console', 'TESTE INTERNO'],
]

table = doc.add_table(rows=len(infra), cols=4)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(infra):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('Custos Mensais:', level=2)
custos = [
    ['Serviço', 'Plano', 'Custo'],
    ['MongoDB Atlas', 'M0 Free (512MB)', 'R$ 0'],
    ['Railway', 'Starter', '~R$ 25/mês (US$ 5)'],
    ['Expo EAS', 'Free', 'R$ 0 (até 30 builds/mês)'],
    ['Google Play', 'Developer', 'R$ 125 (único)'],
    ['Stripe', 'Pay-as-you-go', '3,49% + R$ 0,39 por transação'],
    ['TOTAL FIXO', '', '~R$ 25/mês'],
]

table = doc.add_table(rows=len(custos), cols=3)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(custos):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0 or i == len(custos)-1:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ===== 4. FLUXO DE ATUALIZAÇÕES =====
doc.add_heading('4. Fluxo de Atualizações', level=1)

doc.add_paragraph(
    'IMPORTANTE: O app na Play Store NÃO atualiza automaticamente quando você faz '
    'alterações no código. É necessário seguir o fluxo abaixo para cada atualização:'
)

doc.add_heading('Passo a passo para atualizar o app:', level=2)
steps = [
    'Passo 1 - Fazer alterações: Solicitar as mudanças no Emergent (código é alterado automaticamente)',
    'Passo 2 - Salvar no GitHub: Clicar em "Save to GitHub" no chat do Emergent',
    'Passo 3 - Clonar e buildar: No Prompt de Comando da sua máquina:\n'
    '   cd Desktop\n'
    '   git clone https://github.com/ferreira10cristiano-jpg/itoke-app-final.git itoke-build\n'
    '   cd itoke-build\\frontend\n'
    '   npm install\n'
    '   npx eas-cli build --platform android --profile production --non-interactive',
    'Passo 4 - Upload na Play Store: Baixar o novo AAB do Expo e fazer upload no Google Play Console\n'
    '   Menu: Testar e lançar > Teste interno > Criar nova versão > Enviar AAB',
    'Passo 5 - Promover (quando pronto): Para publicar para todos, clicar em "Promover versão" > Produção',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Quando NÃO precisa rebuildar:', level=2)
doc.add_paragraph('Alterações no backend (servidor) são aplicadas automaticamente no Railway quando você faz Save to GitHub. '
                  'Só precisa rebuildar o AAB quando há mudanças no frontend (telas, layout, funcionalidades do app).')

doc.add_page_break()

# ===== 5. FASE 1 - PRÉ-LANÇAMENTO =====
doc.add_heading('5. Fase 1 - Pré-Lançamento (Semana 1-2)', level=1)

doc.add_heading('5.1 Anti-fraude Básico', level=2)
doc.add_paragraph('Prioridade: ALTA')
items = [
    'Rate Limiting: Limitar tentativas de login (máx. 5 por minuto por IP)',
    'Alertas no Admin: Notificações quando detectar comportamento suspeito',
    'Validação de CPF: Verificar se CPF é válido (algoritmo de dígitos verificadores)',
    'Limite de QR codes por dia por usuário',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2 Histórico de Compras', level=2)
doc.add_paragraph('Prioridade: ALTA')
items = [
    'Tela de histórico de compras de tokens para estabelecimentos',
    'Download de recibo em PDF para cada compra',
    'Filtros por data e status do pagamento',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Testes com Usuários Reais', level=2)
doc.add_paragraph('Prioridade: ALTA')
items = [
    'Convidar 5-10 testadores via link do teste interno da Play Store',
    'Coletar feedback sobre usabilidade',
    'Corrigir bugs encontrados',
    'Testar fluxo completo: cadastro > compra tokens > gerar QR > resgatar oferta',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.4 Configurar Stripe para Produção', level=2)
doc.add_paragraph('Prioridade: ALTA')
items = [
    'Criar conta Stripe em stripe.com (se ainda não criou)',
    'Verificar dados bancários para receber pagamentos',
    'Trocar chave de teste (sk_test_) por chave de produção (sk_live_)',
    'Testar um pagamento real com valor baixo (R$ 1,00)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 6. FASE 2 - LANÇAMENTO =====
doc.add_heading('6. Fase 2 - Lançamento (Semana 3-4)', level=1)

doc.add_heading('6.1 Ficha da Play Store', level=2)
doc.add_paragraph('Prioridade: ALTA (obrigatório para publicação pública)')
items = [
    'Ícone do app (512x512 pixels) - JÁ GERADO',
    'Feature graphic (1024x500 pixels) - banner de destaque',
    'Screenshots do app (mínimo 2, recomendado 8)',
    'Descrição curta (até 80 caracteres): "Ofertas que saem de Graça! Ganhe descontos indicando amigos."',
    'Descrição completa (até 4000 caracteres) - JÁ PREPARADA em DESCRICAO_LOJAS.md',
    'Categoria: Compras',
    'Classificação de conteúdo: Preencher questionário do Google',
    'Informações de contato: email de suporte',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 Promover para Produção', level=2)
items = [
    'No Google Play Console: Teste interno > versão > "Promover versão" > Produção',
    'Aguardar revisão do Google (1-7 dias para produção)',
    'Se rejeitado: ler o motivo, corrigir e resubmeter',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.3 Domínio Customizado', level=2)
items = [
    'Registrar domínio itoke.com.br (se ainda não registrou)',
    'Configurar api.itoke.com.br apontando para o Railway',
    'Criar landing page em itoke.com.br',
    'Atualizar URLs no app para usar o domínio customizado',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 7. FASE 3 - PÓS-LANÇAMENTO =====
doc.add_heading('7. Fase 3 - Pós-Lançamento (Mês 2-3)', level=1)

doc.add_heading('7.1 Google OAuth', level=2)
doc.add_paragraph('Prioridade: MÉDIA')
doc.add_paragraph('Permitir login com conta Google para facilitar o cadastro de novos usuários.')

doc.add_heading('7.2 Refatoração do Backend', level=2)
doc.add_paragraph('Prioridade: MÉDIA')
doc.add_paragraph(
    'O arquivo server.py possui mais de 5000 linhas e precisa ser dividido em módulos:\n'
    '- routes/ (rotas organizadas por domínio)\n'
    '- models/ (modelos de dados)\n'
    '- controllers/ (lógica de negócio)\n'
    'Isso facilita a manutenção e permite que mais desenvolvedores trabalhem simultaneamente.'
)

doc.add_heading('7.3 Build iOS', level=2)
doc.add_paragraph('Prioridade: MÉDIA')
doc.add_paragraph(
    'Para publicar na App Store (iPhone/iPad):\n'
    '- Necessário: Conta Apple Developer (US$ 99/ano = ~R$ 500)\n'
    '- Necessário: Mac para compilação final (pode usar serviço de CI/CD)\n'
    '- O código já está preparado para iOS no eas.json'
)

doc.add_heading('7.4 Segurança Avançada', level=2)
doc.add_paragraph('Prioridade: BAIXA')
items = [
    '2FA (autenticação em dois fatores) para saques de valores altos',
    'Geolocalização para detectar fraudes (ex: QR gerado em cidade diferente do estabelecimento)',
    'Logs de auditoria completos para o Admin',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 8. CRONOGRAMA =====
doc.add_heading('8. Cronograma Detalhado', level=1)

cronograma = [
    ['Semana', 'Tarefa', 'Responsável', 'Status'],
    ['Sem 1', 'Aguardar revisão do Google (teste interno)', 'Google', 'Em andamento'],
    ['Sem 1', 'Testar app no celular via teste interno', 'Você', 'Pendente'],
    ['Sem 1', 'Anti-fraude básico (rate limiting)', 'Emergent', 'Pendente'],
    ['Sem 1', 'Histórico de compras com recibo PDF', 'Emergent', 'Pendente'],
    ['Sem 2', 'Corrigir bugs encontrados nos testes', 'Emergent', 'Pendente'],
    ['Sem 2', 'Configurar Stripe produção', 'Você + Emergent', 'Pendente'],
    ['Sem 2', 'Preparar screenshots e ficha da loja', 'Você + Emergent', 'Pendente'],
    ['Sem 3', 'Promover para Produção na Play Store', 'Você', 'Pendente'],
    ['Sem 3', 'Configurar domínio itoke.com.br', 'Você', 'Pendente'],
    ['Sem 4', 'Landing page itoke.com.br', 'Emergent', 'Pendente'],
    ['Mês 2', 'Google OAuth', 'Emergent', 'Futuro'],
    ['Mês 2', 'Refatoração do backend', 'Emergent', 'Futuro'],
    ['Mês 3', 'Build iOS + App Store', 'Emergent', 'Futuro'],
    ['Mês 3', 'Segurança avançada (2FA, geo)', 'Emergent', 'Futuro'],
]

table = doc.add_table(rows=len(cronograma), cols=4)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(cronograma):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ===== 9. CUSTOS =====
doc.add_heading('9. Custos Operacionais Detalhados', level=1)

doc.add_heading('Custos Fixos (mensais):', level=2)
custos_fix = [
    ['Item', 'Custo Mensal', 'Observação'],
    ['Railway (Backend)', '~R$ 25', 'Plano Starter, escala conforme uso'],
    ['MongoDB Atlas', 'R$ 0', 'Plano gratuito até 512MB'],
    ['Expo EAS', 'R$ 0', 'Até 30 builds/mês grátis'],
    ['TOTAL', '~R$ 25/mês', ''],
]
table = doc.add_table(rows=len(custos_fix), cols=3)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(custos_fix):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0 or i == len(custos_fix)-1:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('Custos Variáveis (por transação):', level=2)
custos_var = [
    ['Item', 'Custo', 'Quando'],
    ['Stripe', '3,49% + R$ 0,39', 'Por cada pagamento recebido'],
    ['Google Play', '15% (primeiro R$ 5M)', 'Por cada compra in-app (se aplicável)'],
]
table = doc.add_table(rows=len(custos_var), cols=3)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(custos_var):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('Custos Únicos (já pagos):', level=2)
custos_uni = [
    ['Item', 'Custo', 'Status'],
    ['Google Play Developer', 'R$ 125', 'PAGO'],
    ['Domínio itoke.com.br', '~R$ 40/ano', 'Pendente'],
    ['Apple Developer (iOS)', '~R$ 500/ano', 'Futuro'],
]
table = doc.add_table(rows=len(custos_uni), cols=3)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(custos_uni):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ===== 10. CREDENCIAIS =====
doc.add_heading('10. Credenciais e Acessos', level=1)
doc.add_paragraph('ATENÇÃO: Guarde este documento em local seguro. Contém credenciais sensíveis.', style='Intense Quote')

creds = [
    ['Serviço', 'Usuário', 'Observação'],
    ['MongoDB Atlas', 'ferreira10cristiano_db_user', 'Senha configurada no Atlas'],
    ['Railway', 'Login via GitHub', 'railway.app'],
    ['Expo/EAS', 'itokecris', 'expo.dev'],
    ['GitHub', 'ferreira10cristiano-jpg', 'github.com'],
    ['Google Play', 'ferreira10cristiano@gmail.com', 'play.google.com/console'],
    ['Stripe', 'A configurar', 'stripe.com'],
]
table = doc.add_table(rows=len(creds), cols=3)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(creds):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('URLs Importantes:', level=2)
urls = [
    'Backend API: https://itoke-app-final-production.up.railway.app',
    'Health Check: https://itoke-app-final-production.up.railway.app/api/health',
    'API Docs: https://itoke-app-final-production.up.railway.app/docs',
    'GitHub: https://github.com/ferreira10cristiano-jpg/itoke-app-final',
    'Expo Builds: https://expo.dev/accounts/itokecris/projects/itoke',
    'MongoDB Atlas: https://cloud.mongodb.com',
    'Railway: https://railway.app',
    'Google Play Console: https://play.google.com/console',
]
for url in urls:
    doc.add_paragraph(url, style='List Bullet')

doc.add_page_break()

# ===== 11. CHECKLIST =====
doc.add_heading('11. Checklist de Lançamento', level=1)

doc.add_heading('Antes de lançar para o público:', level=2)
checks = [
    '[  ] Revisão do Google aprovada (teste interno)',
    '[  ] Stripe configurado para produção (chave sk_live_)',
    '[  ] Testado fluxo completo com pagamento real',
    '[  ] Screenshots do app capturadas (mín. 2)',
    '[  ] Descrição da loja preenchida',
    '[  ] Classificação de conteúdo respondida',
    '[  ] Política de privacidade acessível (URL configurada)',
    '[  ] Ícone e feature graphic configurados',
    '[  ] Email de suporte definido',
    '[  ] Anti-fraude básico implementado',
    '[  ] Testado com pelo menos 5 usuários reais',
    '[  ] Domínio itoke.com.br configurado (opcional)',
    '[  ] Landing page criada (opcional)',
]
for check in checks:
    doc.add_paragraph(check)

doc.add_page_break()

# ===== 12. CONTATOS =====
doc.add_heading('12. Contatos e Suporte', level=1)

doc.add_heading('Desenvolvimento:', level=2)
doc.add_paragraph('Plataforma Emergent - Para qualquer alteração no código, novas funcionalidades ou correção de bugs, '
                  'acesse o chat do Emergent e solicite as mudanças.')

doc.add_heading('Suporte dos Serviços:', level=2)
suporte = [
    ['Serviço', 'Suporte'],
    ['Railway', 'docs.railway.app / Discord do Railway'],
    ['MongoDB Atlas', 'docs.atlas.mongodb.com'],
    ['Expo/EAS', 'docs.expo.dev / Discord do Expo'],
    ['Google Play', 'support.google.com/googleplay/android-developer'],
    ['Stripe', 'stripe.com/docs / Dashboard do Stripe'],
]
table = doc.add_table(rows=len(suporte), cols=2)
table.style = 'Light Grid Accent 1'
for i, row in enumerate(suporte):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell
        if i == 0:
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

# Footer
doc.add_paragraph('')
doc.add_paragraph('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('iToke - Plano de Implantação v1.0 - Abril 2026')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 116, 139)

# Save
output_path = '/app/frontend/iToke_Plano_Implantacao.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
