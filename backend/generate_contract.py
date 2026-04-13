"""
Script to generate the iToke Intermediation Contract in Word format.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(30, 41, 59)

# ===================== HEADER =====================
doc.add_paragraph()
title = doc.add_heading('', level=0)
run = title.add_run('CONTRATO DE INTERMEDIACAO DIGITAL')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(16, 185, 129)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('ENTRE A PLATAFORMA iTOKE E O ESTABELECIMENTO PARCEIRO')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 116, 139)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versao 1.0 — Abril 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(148, 163, 184)
run.italic = True

doc.add_page_break()

# ===================== PARTES =====================
doc.add_heading('DAS PARTES', level=1)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('CONTRATANTE (PLATAFORMA): ')
run.bold = True
p.add_run('iTOKE TECNOLOGIA LTDA, pessoa juridica de direito privado, inscrita no CNPJ sob o n. _________________, '
          'com sede na cidade de _________________, Estado _________________, neste ato representada por seu socio-administrador, '
          'doravante denominada simplesmente "iTOKE" ou "PLATAFORMA".')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('CONTRATADO (ESTABELECIMENTO): ')
run.bold = True
p.add_run('_________________________, pessoa juridica de direito privado, inscrita no CNPJ sob o n. _________________, '
          'com sede na cidade de _________________, Estado _________________, neste ato representada por seu representante legal, '
          'doravante denominada simplesmente "ESTABELECIMENTO" ou "PARCEIRO".')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('As partes acima qualificadas celebram o presente Contrato de Intermediacao Digital, que se regera pelas clausulas e condicoes a seguir:')

doc.add_page_break()

# ===================== CLAUSULA 1 =====================
doc.add_heading('CLAUSULA 1a — DO OBJETO', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '1.1. O presente contrato tem por objeto a prestacao de servicos de intermediacao digital pela PLATAFORMA '
    'ao ESTABELECIMENTO, por meio do aplicativo e plataforma web "iToke", consistindo na disponibilizacao de '
    'ferramentas tecnologicas para que o ESTABELECIMENTO possa criar, publicar e gerenciar ofertas comerciais '
    'com descontos direcionadas a consumidores finais (clientes).'
)
doc.add_paragraph(
    '1.2. A PLATAFORMA atua exclusivamente como INTERMEDIADORA TECNOLOGICA E COMERCIAL entre o ESTABELECIMENTO '
    'e os consumidores finais, NAO se configurando como fornecedora, distribuidora, revendedora ou '
    'representante comercial do ESTABELECIMENTO.'
)
doc.add_paragraph(
    '1.3. Os produtos e servicos ofertados pelo ESTABELECIMENTO por meio da PLATAFORMA sao de inteira '
    'responsabilidade do ESTABELECIMENTO, incluindo sua qualidade, disponibilidade, entrega e conformidade '
    'com a legislacao vigente.'
)

# ===================== CLAUSULA 2 =====================
doc.add_heading('CLAUSULA 2a — DO SISTEMA DE TOKENS E FUNCIONAMENTO', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '2.1. A PLATAFORMA opera por meio de um sistema de "tokens" (creditos digitais) que funcionam como unidade '
    'de acesso as ofertas. Os tokens NAO possuem valor monetario autonomo e NAO constituem moeda eletronica, '
    'criptoativo ou titulo de credito.'
)
doc.add_paragraph(
    '2.2. O ESTABELECIMENTO adquire pacotes de tokens junto a PLATAFORMA, mediante pagamento via cartao de '
    'credito/debito processado pelo sistema Stripe (processador de pagamentos de terceiros).'
)
doc.add_paragraph(
    '2.3. O preco dos pacotes de tokens e definido pela PLATAFORMA e divulgado no aplicativo, podendo ser '
    'alterado a qualquer tempo mediante aviso previo de 15 (quinze) dias.'
)
doc.add_paragraph(
    '2.4. Os tokens sao utilizados pelo ESTABELECIMENTO para ativar ofertas na plataforma, tornando-as visiveis '
    'aos consumidores finais. Cada oferta consome uma quantidade pre-determinada de tokens.'
)
doc.add_paragraph(
    '2.5. Os consumidores finais (clientes) geram QR Codes para resgatar ofertas. O ESTABELECIMENTO valida o '
    'QR Code no momento da compra presencial, confirmando a transacao.'
)

# ===================== CLAUSULA 3 =====================
doc.add_heading('CLAUSULA 3a — DA REMUNERACAO E MODELO FINANCEIRO', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '3.1. A remuneracao da PLATAFORMA consiste na margem retida sobre a venda de pacotes de tokens, '
    'correspondendo a diferenca entre o valor pago pelo comprador e o valor repassavel ao ESTABELECIMENTO.'
)
doc.add_paragraph(
    '3.2. Composicao financeira de cada transacao de venda de tokens:'
)
doc.add_paragraph()

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ('Valor pago pelo comprador (ex: R$ 9,90)', '100% do valor'),
    ('Taxa do processador de pagamento (Stripe)', 'Variavel (~3,99% + R$ 0,39)'),
    ('Margem da PLATAFORMA (receita de intermediacao)', 'Ex: R$ 6,90 por pacote'),
    ('Valor repassavel ao ESTABELECIMENTO (creditos)', 'Ex: R$ 3,00 por pacote'),
]
table.rows[0].cells[0].text = 'Componente'
table.rows[0].cells[1].text = 'Valor'
for i, (a, b) in enumerate(data):
    table.rows[i+1].cells[0].text = a
    table.rows[i+1].cells[1].text = b

doc.add_paragraph()
doc.add_paragraph(
    '3.3. O valor repassavel ao ESTABELECIMENTO e acumulado na forma de creditos digitais dentro da '
    'PLATAFORMA, oriundos das transacoes realizadas por consumidores finais que utilizaram QR Codes nas '
    'ofertas do ESTABELECIMENTO.'
)
doc.add_paragraph(
    '3.4. Os creditos acumulados pelo ESTABELECIMENTO constituem REPASSE de valores, e NAO receita da '
    'PLATAFORMA. A PLATAFORMA emitira Nota Fiscal de Servico Eletronica (NFS-e) exclusivamente sobre sua '
    'margem de intermediacao (conforme item 3.2), e nao sobre o valor total das transacoes.'
)
doc.add_paragraph(
    '3.5. Os valores dos pacotes e a composicao financeira poderao ser reajustados pela PLATAFORMA mediante '
    'comunicacao previa de 30 (trinta) dias ao ESTABELECIMENTO.'
)

# ===================== CLAUSULA 4 =====================
doc.add_heading('CLAUSULA 4a — DOS SAQUES E REPASSES', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '4.1. O ESTABELECIMENTO podera solicitar o saque dos creditos acumulados a qualquer momento, desde que '
    'o saldo minimo seja de R$ 10,00 (dez reais) ou outro valor definido pela PLATAFORMA.'
)
doc.add_paragraph(
    '4.2. Os saques serao processados via transferencia PIX para a conta bancaria informada pelo '
    'ESTABELECIMENTO, no prazo de ate 5 (cinco) dias uteis apos a aprovacao pela PLATAFORMA.'
)
doc.add_paragraph(
    '4.3. A PLATAFORMA reserva-se o direito de reter ou suspender saques em caso de suspeita de fraude, '
    'irregularidade ou violacao deste contrato, pelo tempo necessario para apuracao.'
)
doc.add_paragraph(
    '4.4. E de inteira responsabilidade do ESTABELECIMENTO informar e manter atualizados seus dados '
    'bancarios na plataforma. A PLATAFORMA nao se responsabiliza por transferencias realizadas para contas '
    'informadas incorretamente pelo ESTABELECIMENTO.'
)

# ===================== CLAUSULA 5 =====================
doc.add_heading('CLAUSULA 5a — DAS OBRIGACOES DO ESTABELECIMENTO', level=1)
doc.add_paragraph()
doc.add_paragraph('5.1. O ESTABELECIMENTO se obriga a:')
doc.add_paragraph('a) Manter informacoes verdadeiras, completas e atualizadas na plataforma, incluindo razao social, CNPJ, endereco e dados de contato;', style='List Bullet')
doc.add_paragraph('b) Criar ofertas que correspondam fielmente aos produtos e servicos disponiveis em seu estabelecimento fisico;', style='List Bullet')
doc.add_paragraph('c) Honrar todas as ofertas publicadas na plataforma, pelo preco e condicoes anunciados, durante o periodo de vigencia;', style='List Bullet')
doc.add_paragraph('d) Validar os QR Codes apresentados pelos consumidores de forma diligente e honesta;', style='List Bullet')
doc.add_paragraph('e) Atender os consumidores com respeito e qualidade, sem discriminacao entre clientes provenientes do iToke e demais clientes;', style='List Bullet')
doc.add_paragraph('f) Cumprir toda a legislacao aplicavel, incluindo o Codigo de Defesa do Consumidor (Lei 8.078/1990), normas sanitarias e tributarias;', style='List Bullet')
doc.add_paragraph('g) Emitir nota fiscal ou cupom fiscal ao consumidor final pela venda realizada, quando exigido pela legislacao;', style='List Bullet')
doc.add_paragraph('h) Nao utilizar a plataforma para atividades ilicitas, fraudulentas ou que violem direitos de terceiros;', style='List Bullet')
doc.add_paragraph('i) Comunicar imediatamente a PLATAFORMA sobre qualquer irregularidade, reclamacao ou problema relacionado as ofertas.', style='List Bullet')

# ===================== CLAUSULA 6 =====================
doc.add_heading('CLAUSULA 6a — DAS OBRIGACOES DA PLATAFORMA', level=1)
doc.add_paragraph()
doc.add_paragraph('6.1. A PLATAFORMA se obriga a:')
doc.add_paragraph('a) Disponibilizar e manter o aplicativo e plataforma web em funcionamento, com nivel de disponibilidade razoavel, ressalvadas manutencoes programadas e eventos de forca maior;', style='List Bullet')
doc.add_paragraph('b) Processar os pagamentos de forma segura, por meio de processador de pagamentos certificado (Stripe);', style='List Bullet')
doc.add_paragraph('c) Efetuar os repasses de creditos ao ESTABELECIMENTO conforme as regras estabelecidas neste contrato;', style='List Bullet')
doc.add_paragraph('d) Proteger os dados pessoais dos usuarios conforme a Lei Geral de Protecao de Dados (LGPD — Lei 13.709/2018);', style='List Bullet')
doc.add_paragraph('e) Fornecer ao ESTABELECIMENTO acesso a dashboard com metricas de desempenho (visualizacoes, QR Codes gerados, vendas, creditos);', style='List Bullet')
doc.add_paragraph('f) Disponibilizar relatorios fiscais para fins contabeis;', style='List Bullet')
doc.add_paragraph('g) Emitir NFS-e sobre a receita de intermediacao (sua margem) conforme legislacao vigente.', style='List Bullet')

# ===================== CLAUSULA 7 =====================
doc.add_heading('CLAUSULA 7a — DA PROTECAO DE DADOS (LGPD)', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '7.1. As partes se comprometem a cumprir integralmente a Lei Geral de Protecao de Dados Pessoais '
    '(Lei 13.709/2018 — LGPD) no tratamento de dados pessoais dos consumidores e de qualquer outra '
    'pessoa natural envolvida na operacao.'
)
doc.add_paragraph(
    '7.2. A PLATAFORMA atuara como CONTROLADORA dos dados pessoais coletados dos consumidores no ambito '
    'do aplicativo, sendo responsavel pela definicao das finalidades e meios de tratamento.'
)
doc.add_paragraph(
    '7.3. O ESTABELECIMENTO tera acesso apenas a dados estritamente necessarios para a operacao, tais como: '
    'nome do consumidor (quando aplicavel), CPF parcialmente mascarado (ex: 123.***.***-01), valor da transacao '
    'e identificador do QR Code. O ESTABELECIMENTO NAO tera acesso a dados completos de CPF, e-mail ou '
    'telefone dos consumidores.'
)
doc.add_paragraph(
    '7.4. O ESTABELECIMENTO compromete-se a NAO coletar, armazenar, compartilhar ou utilizar dados pessoais '
    'dos consumidores obtidos por meio da plataforma para qualquer finalidade diversa da execucao da transacao.'
)
doc.add_paragraph(
    '7.5. Em caso de incidente de seguranca envolvendo dados pessoais, ambas as partes se comprometem a '
    'comunicar mutuamente no prazo de 48 (quarenta e oito) horas, alem de adotar as providencias legais '
    'perante a Autoridade Nacional de Protecao de Dados (ANPD).'
)

# ===================== CLAUSULA 8 =====================
doc.add_heading('CLAUSULA 8a — DA PROPRIEDADE INTELECTUAL', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '8.1. A marca "iToke", o aplicativo, o codigo-fonte, o design, os algoritmos e toda a tecnologia da '
    'PLATAFORMA sao de propriedade exclusiva da iTOKE TECNOLOGIA LTDA, protegidos pela Lei de Propriedade '
    'Industrial (Lei 9.279/1996) e pela Lei de Direitos Autorais (Lei 9.610/1998).'
)
doc.add_paragraph(
    '8.2. O presente contrato NAO confere ao ESTABELECIMENTO qualquer direito de propriedade intelectual '
    'sobre a PLATAFORMA. O ESTABELECIMENTO recebe apenas uma licenca limitada, nao exclusiva e revogavel '
    'para utilizar os servicos durante a vigencia deste contrato.'
)
doc.add_paragraph(
    '8.3. O ESTABELECIMENTO autoriza a PLATAFORMA a exibir seu nome comercial, logotipo e endereco no '
    'aplicativo para fins de divulgacao das ofertas, sem que isso implique em custos adicionais ou '
    'transferencia de direitos.'
)

# ===================== CLAUSULA 9 =====================
doc.add_heading('CLAUSULA 9a — DA RESPONSABILIDADE', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '9.1. A PLATAFORMA NAO se responsabiliza por:'
)
doc.add_paragraph('a) A qualidade, seguranca ou legalidade dos produtos e servicos ofertados pelo ESTABELECIMENTO;', style='List Bullet')
doc.add_paragraph('b) O cumprimento das ofertas pelo ESTABELECIMENTO perante os consumidores;', style='List Bullet')
doc.add_paragraph('c) Danos causados ao consumidor em decorrencia de acao ou omissao do ESTABELECIMENTO;', style='List Bullet')
doc.add_paragraph('d) Indisponibilidade do servico por caso fortuito, forca maior ou manutencao programada;', style='List Bullet')
doc.add_paragraph('e) Perdas e danos decorrentes de uso indevido da plataforma pelo ESTABELECIMENTO.', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    '9.2. O ESTABELECIMENTO sera o unico responsavel perante os orgaos de defesa do consumidor e o Poder '
    'Judiciario por reclamacoes, devolucoes e indenizacoes relacionadas aos seus produtos e servicos, '
    'devendo isentar a PLATAFORMA de qualquer responsabilidade nesse sentido.'
)
doc.add_paragraph(
    '9.3. Nao obstante o disposto acima, as partes reconhecem que, conforme jurisprudencia vigente, '
    'a PLATAFORMA podera ser responsabilizada solidariamente em determinadas situacoes previstas no '
    'Codigo de Defesa do Consumidor, cabendo ao ESTABELECIMENTO o ressarcimento integral de valores '
    'eventualmente pagos pela PLATAFORMA nessas hipoteses.'
)

# ===================== CLAUSULA 10 =====================
doc.add_heading('CLAUSULA 10a — DA VIGENCIA E RESCISAO', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '10.1. O presente contrato entra em vigor na data de sua assinatura (fisica ou digital) e tem prazo '
    'de vigencia indeterminado.'
)
doc.add_paragraph(
    '10.2. Qualquer das partes podera rescindir este contrato a qualquer tempo, mediante comunicacao '
    'por escrito (incluindo e-mail) com antecedencia minima de 30 (trinta) dias.'
)
doc.add_paragraph(
    '10.3. Em caso de rescisao, a PLATAFORMA efetuara o repasse dos creditos acumulados pelo '
    'ESTABELECIMENTO no prazo de 15 (quinze) dias uteis apos a data de rescisao, descontadas eventuais '
    'pendencias, multas ou ressarcimentos devidos.'
)
doc.add_paragraph(
    '10.4. A PLATAFORMA podera rescindir imediatamente este contrato, sem aviso previo, em caso de:'
)
doc.add_paragraph('a) Fraude ou tentativa de fraude pelo ESTABELECIMENTO;', style='List Bullet')
doc.add_paragraph('b) Violacao de qualquer clausula deste contrato;', style='List Bullet')
doc.add_paragraph('c) Conduta que prejudique a imagem ou reputacao da PLATAFORMA;', style='List Bullet')
doc.add_paragraph('d) Inatividade do ESTABELECIMENTO por periodo superior a 90 (noventa) dias consecutivos;', style='List Bullet')
doc.add_paragraph('e) Determinacao judicial ou administrativa.', style='List Bullet')

# ===================== CLAUSULA 11 =====================
doc.add_heading('CLAUSULA 11a — DO SISTEMA ANTI-FRAUDE', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '11.1. A PLATAFORMA mantem sistema anti-fraude ativo, incluindo:'
)
doc.add_paragraph('a) Limitacao de taxa (rate limiting) para operacoes criticas;', style='List Bullet')
doc.add_paragraph('b) Validacao de CPF com algoritmo verificador de digitos;', style='List Bullet')
doc.add_paragraph('c) Validacao de CNPJ;', style='List Bullet')
doc.add_paragraph('d) Deteccao de contas duplicadas e auto-indicacao;', style='List Bullet')
doc.add_paragraph('e) Monitoramento de atividades suspeitas com alertas automaticos.', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    '11.2. O ESTABELECIMENTO compromete-se a cooperar com investigacoes de fraude, fornecendo informacoes '
    'e documentos solicitados pela PLATAFORMA no prazo de 5 (cinco) dias uteis.'
)

# ===================== CLAUSULA 12 =====================
doc.add_heading('CLAUSULA 12a — DOS TOKENS GRATUITOS DE REPRESENTANTES', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '12.1. A PLATAFORMA podera, a seu criterio, disponibilizar tokens gratuitos ao ESTABELECIMENTO por '
    'intermedio de representantes comerciais parceiros, como incentivo para adesao a plataforma.'
)
doc.add_paragraph(
    '12.2. Os tokens gratuitos possuem prazo de validade definido pela PLATAFORMA (padrao: 30 dias), '
    'apos o qual serao automaticamente cancelados se nao utilizados.'
)
doc.add_paragraph(
    '12.3. Tokens gratuitos NAO geram creditos de repasse ao ESTABELECIMENTO. Sao exclusivamente para '
    'permitir a publicacao de ofertas iniciais como demonstracao da plataforma.'
)
doc.add_paragraph(
    '12.4. A concessao de tokens gratuitos nao implica em obrigacao de concessoes futuras ou em direito '
    'adquirido do ESTABELECIMENTO.'
)

# ===================== CLAUSULA 13 =====================
doc.add_heading('CLAUSULA 13a — DISPOSICOES GERAIS', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '13.1. Este contrato constitui o acordo integral entre as partes sobre o objeto descrito, substituindo '
    'quaisquer entendimentos ou acordos anteriores, verbais ou escritos.'
)
doc.add_paragraph(
    '13.2. A tolerancia de qualquer das partes em relacao ao descumprimento de clausulas deste contrato '
    'nao constituira novacao, renuncia ou precedente invocavel pela outra parte.'
)
doc.add_paragraph(
    '13.3. Se qualquer clausula deste contrato for considerada invalida ou inexequivel, as demais clausulas '
    'permanecerao em pleno vigor e efeito.'
)
doc.add_paragraph(
    '13.4. A PLATAFORMA podera alterar os termos deste contrato mediante comunicacao previa de 30 (trinta) '
    'dias ao ESTABELECIMENTO. A continuidade do uso da plataforma apos o prazo constituira aceite tacito das '
    'alteracoes.'
)
doc.add_paragraph(
    '13.5. O presente contrato podera ser aceito de forma digital (aceite eletronico) no momento do cadastro '
    'do ESTABELECIMENTO na plataforma, tendo a mesma validade juridica da assinatura fisica, nos termos da '
    'Medida Provisoria 2.200-2/2001 e do artigo 10 da Lei 12.965/2014 (Marco Civil da Internet).'
)

# ===================== CLAUSULA 14 =====================
doc.add_heading('CLAUSULA 14a — DO FORO', level=1)
doc.add_paragraph()
doc.add_paragraph(
    '14.1. As partes elegem o foro da Comarca de _________________, Estado de _________________, para '
    'dirimir quaisquer controversias decorrentes deste contrato, renunciando a qualquer outro, por mais '
    'privilegiado que seja.'
)

doc.add_paragraph()
doc.add_paragraph()

# ===================== SIGNATURES =====================
doc.add_heading('ASSINATURAS', level=1)
doc.add_paragraph()
doc.add_paragraph('_________________, _____ de _________________ de 20_____.')
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('_____________________________________________')
p = doc.add_paragraph()
run = p.add_run('iTOKE TECNOLOGIA LTDA')
run.bold = True
doc.add_paragraph('CNPJ: _________________')
doc.add_paragraph('Representante Legal: _________________')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('_____________________________________________')
p = doc.add_paragraph()
run = p.add_run('ESTABELECIMENTO PARCEIRO')
run.bold = True
doc.add_paragraph('CNPJ: _________________')
doc.add_paragraph('Razao Social: _________________')
doc.add_paragraph('Representante Legal: _________________')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TESTEMUNHAS')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('1. Nome: _________________________ CPF: _________________')
doc.add_paragraph()
doc.add_paragraph('2. Nome: _________________________ CPF: _________________')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento gerado pela plataforma iToke — Versao 1.0')
run.italic = True
run.font.color.rgb = RGBColor(148, 163, 184)
run.font.size = Pt(9)

output_path = '/app/frontend/iToke_Contrato_Intermediacao_Estabelecimento.docx'
doc.save(output_path)
print(f'Contract saved to {output_path}')
